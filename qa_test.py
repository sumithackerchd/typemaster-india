"""Comprehensive QA test suite for TypeMaster India admin + features."""

import io
import json
import os
import sys
import tempfile
from datetime import datetime

from werkzeug.security import generate_password_hash

from app import app
from models import db
from models.user import User
from models.result import Result
from utils.extract import ocr_available

PASS = 0
FAIL = 0
WARN = 0
RESULTS = []


def ok(name, detail=""):
    global PASS
    PASS += 1
    RESULTS.append(("PASS", name, detail))
    print(f"  PASS  {name}" + (f" — {detail}" if detail else ""))


def fail(name, detail=""):
    global FAIL
    FAIL += 1
    RESULTS.append(("FAIL", name, detail))
    print(f"  FAIL  {name}" + (f" — {detail}" if detail else ""))


def warn(name, detail=""):
    global WARN
    WARN += 1
    RESULTS.append(("WARN", name, detail))
    print(f"  WARN  {name}" + (f" — {detail}" if detail else ""))


def login_as(client, user):
    with client.session_transaction() as sess:
        sess["_user_id"] = str(user.id)
        sess["_fresh"] = True


# ---------------------------------------------------------------------------
SIDEBAR_ROUTES = [
    ("/admin/", "Dashboard"),
    ("/admin/users", "Users"),
    ("/admin/results", "Results"),
    ("/admin/paragraphs/english", "English Paragraphs"),
    ("/admin/paragraphs/hindi", "Hindi Paragraphs"),
    ("/custom/", "Custom Practice"),
    ("/admin/downloads", "Downloads Manager"),
    ("/admin/settings", "Settings"),
]

PUBLIC_ROUTES = [
    "/", "/about", "/privacy", "/terms", "/contact",
    "/login", "/register", "/downloads",
    "/dashboard", "/typing", "/hindi", "/history",
    "/leaderboard", "/profile", "/mock-test", "/result",
]

with app.app_context():
    db.create_all()
    admin = User.query.filter_by(username="qaadmin").first()
    if not admin:
        admin = User(
            full_name="QA Admin",
            username="qaadmin",
            email="qa@test.com",
            mobile="8888888888",
            password=generate_password_hash("test1234"),
            is_admin=True,
        )
        db.session.add(admin)
        db.session.commit()
    elif not admin.is_admin:
        admin.is_admin = True
        db.session.commit()

    admin_id = admin.id

    result_row = Result.query.filter_by(user_id=admin_id).first()
    if not result_row:
        result_row = Result(
            user_id=admin_id,
            language="English",
            difficulty="Easy",
            duration=60,
            gross_wpm=45,
            net_wpm=42,
            cpm=210,
            accuracy=98.0,
            errors=2,
            created_at=datetime.utcnow(),
        )
        db.session.add(result_row)
        db.session.commit()

    result_id = result_row.id

client = app.test_client()
with app.app_context():
    admin = db.session.get(User, admin_id)
    login_as(client, admin)

print("\n=== 1. SIDEBAR ROUTES ===")
for route, label in SIDEBAR_ROUTES:
    resp = client.get(route, follow_redirects=False)
    is_admin_route = route.startswith("/admin")
    if resp.status_code == 200 and (not is_admin_route or b"admin-sidebar" in resp.data):
        ok(f"Sidebar: {label}", route)
    else:
        fail(f"Sidebar: {label}", f"{route} -> {resp.status_code}")

print("\n=== 2. ADMIN PAGES ===")
for route, label in SIDEBAR_ROUTES:
    resp = client.get(route, follow_redirects=False)
    body = resp.get_data(as_text=True)
    if resp.status_code != 200:
        fail(f"Admin page {label}", str(resp.status_code))
    elif "Traceback" in body or "Internal Server Error" in body:
        fail(f"Admin page {label}", "server error in body")
    elif route.startswith("/admin") and "admin-sidebar" not in body:
        fail(f"Admin page {label}", "not admin layout")
    else:
        ok(f"Admin page {label}")

print("\n=== 3. CRUD OPERATIONS ===")
# Add
resp = client.post(
    "/admin/paragraphs/save",
    json={"language": "english", "difficulty": "easy", "action": "add", "content": "QA test paragraph one."},
)
data = resp.get_json()
if data and data.get("success"):
    ok("CRUD: Add paragraph")
    idx = len(data["paragraphs"]["easy"]) - 1
else:
    fail("CRUD: Add paragraph", str(data))
    idx = 0

# Edit
resp = client.post(
    "/admin/paragraphs/save",
    json={"language": "english", "difficulty": "easy", "action": "edit", "index": idx, "content": "QA test paragraph edited."},
)
data = resp.get_json()
if data and data.get("success") and "edited" in data["paragraphs"]["easy"][idx]:
    ok("CRUD: Edit paragraph")
else:
    fail("CRUD: Edit paragraph", str(data))

# Delete
resp = client.post(
    "/admin/paragraphs/save",
    json={"language": "english", "difficulty": "easy", "action": "delete", "index": idx},
)
data = resp.get_json()
if data and data.get("success"):
    ok("CRUD: Delete paragraph")
else:
    fail("CRUD: Delete paragraph", str(data))

# User search
resp = client.get("/admin/users?search=qaadmin")
if b"qaadmin" in resp.data:
    ok("CRUD: User search")
else:
    fail("CRUD: User search")

print("\n=== 4-7. IMPORT TESTS ===")

def make_import(content, filename, lang="english", diff="medium"):
    buf = io.BytesIO(content if isinstance(content, bytes) else content.encode("utf-8"))
    buf.seek(0)
    return client.post(
        "/admin/paragraphs/import",
        data={"file": (buf, filename), "language": lang, "difficulty": diff},
        content_type="multipart/form-data",
    )


# Bulk TXT
r = make_import("First bulk paragraph.\n\nSecond bulk paragraph.", "bulk.txt")
d = r.get_json()
if d and d.get("success") and d.get("added", 0) >= 1:
    ok("Import: Bulk TXT", f"added={d['added']}")
else:
    fail("Import: Bulk TXT", str(d))

# JSON
payload = json.dumps({"english": {"easy": ["JSON import one.", "JSON import two."]}})
r = make_import(payload.encode(), "bulk.json")
d = r.get_json()
if d and d.get("success"):
    ok("Import: JSON")
else:
    fail("Import: JSON", str(d))

# CSV
r = make_import("difficulty,content\neasy,CSV import paragraph.", "bulk.csv")
d = r.get_json()
if d and d.get("success"):
    ok("Import: CSV")
else:
    fail("Import: CSV", str(d))

# DOCX
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("DOCX import paragraph one.")
    doc.add_paragraph("DOCX import paragraph two.")
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    r = client.post(
        "/admin/paragraphs/import",
        data={"file": (bio, "bulk.docx"), "language": "english", "difficulty": "medium"},
        content_type="multipart/form-data",
    )
    d = r.get_json()
    if d and d.get("success"):
        ok("Import: DOCX", f"added={d.get('added')}")
    else:
        fail("Import: DOCX", str(d))
except Exception as exc:
    fail("Import: DOCX", str(exc))

# PDF
try:
    from reportlab.pdfgen import canvas as rl_canvas
    pdf_path = os.path.join(tempfile.gettempdir(), "qa_import.pdf")
    c = rl_canvas.Canvas(pdf_path)
    c.drawString(72, 720, "PDF import paragraph one.")
    c.drawString(72, 700, "PDF import paragraph two.")
    c.save()
    with open(pdf_path, "rb") as f:
        r = client.post(
            "/admin/paragraphs/import",
            data={"file": (f, "bulk.pdf"), "language": "english", "difficulty": "hard"},
            content_type="multipart/form-data",
        )
    d = r.get_json()
    if d and d.get("success"):
        ok("Import: PDF", f"added={d.get('added')}")
    else:
        fail("Import: PDF", str(d))
except Exception as exc:
    fail("Import: PDF", str(exc))

# OCR
if ocr_available():
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new("RGB", (400, 100), "white")
        draw = ImageDraw.Draw(img)
        draw.text((10, 30), "OCR import paragraph text.", fill="black")
        bio = io.BytesIO()
        img.save(bio, format="PNG")
        bio.seek(0)
        r = client.post(
            "/admin/paragraphs/import",
            data={"file": (bio, "ocr.png"), "language": "english", "difficulty": "easy"},
            content_type="multipart/form-data",
        )
        d = r.get_json()
        if d and d.get("success"):
            ok("Import: OCR Image", f"added={d.get('added')}")
        else:
            warn("Import: OCR Image", str(d))
    except Exception as exc:
        warn("Import: OCR Image", str(exc))
else:
    warn("Import: OCR Image", "Tesseract not available on server")

print("\n=== 8-10. EXPORT TESTS ===")
for fmt in ("json", "txt", "csv"):
    resp = client.get(f"/admin/paragraphs/export?language=english&format={fmt}")
    if resp.status_code == 200 and len(resp.data) > 10:
        ok(f"Export: {fmt.upper()}", f"{len(resp.data)} bytes")
    else:
        fail(f"Export: {fmt.upper()}", str(resp.status_code))

print("\n=== 11-13. PAGINATION / SEARCH / FILTERS (HTML checks) ===")
resp = client.get("/admin/paragraphs/english")
body = resp.get_data(as_text=True)
for marker in ("paragraphSearch", "paragraphPager", "difficultySelect"):
    if marker in body:
        ok(f"UI element: {marker}")
    else:
        fail(f"UI element: {marker}")

if "importProgress" in body and "importStatus" in body:
    ok("Upload progress UI present")
else:
    fail("Upload progress UI present")

print("\n=== 14. MOBILE SIDEBAR / RESPONSIVE (markup) ===")
resp = client.get("/admin/")
body = resp.get_data(as_text=True)
for marker in ("adminSidebar", "admin-toggle", "toggleAdminSidebar", "admin-backdrop"):
    if marker in body:
        ok(f"Responsive: {marker}")
    else:
        fail(f"Responsive: {marker}")

print("\n=== 17-20. CERTIFICATE / QR / VERIFY ===")
from utils.certificate import ensure_certificate_identity, verification_url, build_qr_svg_data_uri
from utils.settings_store import signature_abs_path

with app.app_context():
    result_row = db.session.get(Result, result_id)
    cert_id, token = ensure_certificate_identity(result_row)
    verify_url = verification_url("http://localhost:5007/", cert_id)
    qr = build_qr_svg_data_uri(verify_url)

if cert_id and token:
    ok("Certificate ID + token", cert_id)
else:
    fail("Certificate ID + token")

if qr.startswith("data:image/svg"):
    ok("QR code generation")
else:
    fail("QR code generation")

sig = signature_abs_path()
if sig and os.path.exists(sig):
    ok("Signature PNG exists", sig)
else:
    warn("Signature PNG exists", "Place static/images/signature.png for certificates")

resp = client.get(f"/certificate/{result_id}")
if resp.status_code == 200 and cert_id.encode() in resp.data:
    ok("Certificate HTML page")
else:
    fail("Certificate HTML page", str(resp.status_code))

resp = client.get(f"/verify/{cert_id}")
if resp.status_code == 200 and b"QA Admin" in resp.data:
    ok("Certificate verification page")
else:
    fail("Certificate verification page", str(resp.status_code))

resp = client.get(f"/certificate/download/{result_id}")
if resp.status_code == 200 and resp.data[:4] == b"%PDF":
    ok("Certificate PDF download", f"{len(resp.data)} bytes")
else:
    fail("Certificate PDF download", str(resp.status_code))

print("\n=== 21. DOWNLOADS PAGE ===")
resp = client.get("/downloads")
if resp.status_code == 200:
    ok("Public Downloads page")
else:
    fail("Public Downloads page")

resp = client.get("/admin/downloads")
if resp.status_code == 200:
    ok("Admin Downloads page")
else:
    fail("Admin Downloads page")

print("\n=== 22-23. ALL PUBLIC PAGES ===")
for route in PUBLIC_ROUTES:
    resp = client.get(route, follow_redirects=True)
    body = resp.get_data(as_text=True)
    if resp.status_code == 200 and "Traceback" not in body:
        ok(f"Page: {route}")
    elif resp.status_code in (302, 308):
        ok(f"Page: {route} (redirect)")
    else:
        fail(f"Page: {route}", str(resp.status_code))

print("\n=== 24. INVALID VERIFY ===")
resp = client.get("/verify/TMI-FAKE-000000")
if resp.status_code == 200:
    ok("Invalid certificate shows verify page")
else:
    fail("Invalid certificate verify", str(resp.status_code))

print("\n=== 25. NON-ADMIN SECURITY ===")
with app.app_context():
    normal = User.query.filter_by(is_admin=False).first()
    if not normal:
        normal = User(
            full_name="QA User",
            username="qauser",
            email="qauser@test.com",
            mobile="7777777777",
            password=generate_password_hash("test1234"),
            is_admin=False,
        )
        db.session.add(normal)
        db.session.commit()

    normal_id = normal.id

normal_client = app.test_client()
with app.app_context():
    normal = db.session.get(User, normal_id)
    login_as(normal_client, normal)

ADMIN_BLOCKED = [
    "/admin/", "/admin/users", "/admin/results",
    "/admin/settings", "/admin/downloads",
    "/admin/paragraphs/english", "/admin/paragraphs/hindi",
]
for route in ADMIN_BLOCKED:
    resp = normal_client.get(route, follow_redirects=False)
    if resp.status_code in (302, 303):
        ok(f"Non-admin blocked: {route}", "redirect")
    elif resp.status_code == 403:
        ok(f"Non-admin blocked: {route}", "403")
    else:
        fail(f"Non-admin blocked: {route}", f"got {resp.status_code}")

resp = normal_client.get("/admin/", follow_redirects=True)
if b"You are not authorized" in resp.data or b"Access Denied" in resp.data:
    ok("Non-admin sees authorization message")
else:
    fail("Non-admin authorization message")

print("\n" + "=" * 60)
print(f"QA SUMMARY: {PASS} passed, {FAIL} failed, {WARN} warnings")
print("=" * 60)
sys.exit(1 if FAIL else 0)
