"""Text extraction + OCR helpers for the import / custom-practice features.

Supports: TXT, JSON, CSV, XLSX, DOCX, PDF and images (JPG/PNG/JPEG/WEBP).

Image and (image-only) PDF extraction use Tesseract via ``pytesseract``. That
requires the Tesseract binary to be installed on the host. When the binary is
missing (e.g. some sandboxes) OCR raises :class:`OCRUnavailable` with a clear
message; every other format keeps working normally.
"""

import csv
import io
import json
import os
import re


TEXT_EXTS = {".txt"}
JSON_EXTS = {".json"}
CSV_EXTS = {".csv"}
EXCEL_EXTS = {".xlsx", ".xls"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".gif"}

SUPPORTED_EXTS = (
    TEXT_EXTS | JSON_EXTS | CSV_EXTS | EXCEL_EXTS | DOCX_EXTS | PDF_EXTS | IMAGE_EXTS
)


class ExtractionError(Exception):
    """Raised when a file cannot be parsed."""


class OCRUnavailable(ExtractionError):
    """Raised when OCR is requested but Tesseract is not installed."""


# ---------------------------------------------------------------------------
# OCR
# ---------------------------------------------------------------------------

def ocr_available():
    """Return True if pytesseract + a Tesseract binary are usable."""
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except Exception:
        return False

    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _tess_lang(lang):
    """Map our language slugs to Tesseract language codes.

    Defaults to ``hin+eng`` so mixed English + Hindi documents are handled
    well even when the caller doesn't specify a language.
    """
    if lang == "hindi":
        return "hin+eng"
    if lang == "english":
        return "eng"
    return "hin+eng"


def _preprocess_for_ocr(pil_image):
    """Clean an image for OCR: grayscale, denoise, deskew and threshold.

    Uses OpenCV + NumPy when available; falls back to a plain grayscale PIL
    image if OpenCV isn't installed so OCR still works (just less accurately).
    """
    try:
        import cv2
        import numpy as np
        from PIL import Image
    except Exception:  # noqa: BLE001
        # OpenCV unavailable — return a simple grayscale image.
        try:
            return pil_image.convert("L")
        except Exception:  # noqa: BLE001
            return pil_image

    try:
        img = np.array(pil_image.convert("RGB"))
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)

        # Upscale small images for better recognition.
        h, w = gray.shape[:2]
        if max(h, w) < 1000:
            scale = 1000.0 / max(h, w)
            gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

        # Denoise while keeping edges crisp.
        gray = cv2.fastNlMeansDenoising(gray, h=10)

        # Deskew based on the dominant text angle.
        coords = np.column_stack(np.where(gray < 128))
        if coords.size > 0:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            if abs(angle) > 0.5:
                (ih, iw) = gray.shape[:2]
                M = cv2.getRotationMatrix2D((iw // 2, ih // 2), angle, 1.0)
                gray = cv2.warpAffine(
                    gray, M, (iw, ih),
                    flags=cv2.INTER_CUBIC,
                    borderMode=cv2.BORDER_REPLICATE,
                )

        # Adaptive threshold to a clean black-on-white image.
        gray = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 31, 15,
        )

        return Image.fromarray(gray)
    except Exception:  # noqa: BLE001
        try:
            return pil_image.convert("L")
        except Exception:  # noqa: BLE001
            return pil_image


def _ocr_pil(image, lang=None):
    """OCR a PIL image (after preprocessing) with a graceful language fallback."""
    import pytesseract

    processed = _preprocess_for_ocr(image)
    tess_lang = _tess_lang(lang)

    try:
        text = pytesseract.image_to_string(processed, lang=tess_lang)
    except Exception:
        # Language pack missing or preprocessing hurt — retry on the original.
        try:
            text = pytesseract.image_to_string(image, lang=tess_lang)
        except Exception:
            text = pytesseract.image_to_string(image)

    # If preprocessing produced almost nothing, retry on the raw image.
    if len(text.strip()) < 3:
        try:
            raw = pytesseract.image_to_string(image, lang=tess_lang)
            if len(raw.strip()) > len(text.strip()):
                text = raw
        except Exception:
            pass

    return text


def ocr_image_bytes(data, lang=None):
    """Run OCR on raw image bytes and return extracted text."""
    if not ocr_available():
        raise OCRUnavailable(
            "Image text recognition (OCR) isn't available right now. "
            "You can still paste text directly or upload a TXT, DOCX or PDF file."
        )

    from PIL import Image

    try:
        image = Image.open(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(
            "That image couldn't be opened. Please upload a valid PNG, JPG or WEBP file."
        )

    try:
        return _ocr_pil(image, lang=lang)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Text recognition failed on that image: {exc}")


# ---------------------------------------------------------------------------
# Per-format extractors (all return a single text blob)
# ---------------------------------------------------------------------------

def _extract_txt(data):
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_json(data):
    text = _extract_txt(data)
    try:
        obj = json.loads(text)
    except ValueError as exc:
        raise ExtractionError(f"Invalid JSON: {exc}")

    chunks = []

    def walk(node):
        if isinstance(node, str):
            chunks.append(node)
        elif isinstance(node, list):
            for item in node:
                walk(item)
        elif isinstance(node, dict):
            # Prefer common content keys, otherwise walk all values.
            for value in node.values():
                walk(value)

    walk(obj)
    return "\n\n".join(c for c in chunks if c and c.strip())


def _extract_csv(data):
    text = _extract_txt(data)
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        cell = " ".join(c.strip() for c in row if c and c.strip())
        if cell:
            rows.append(cell)
    return "\n".join(rows)


def _extract_excel(data):
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Excel support unavailable: {exc}")

    try:
        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Could not read the Excel file: {exc}")

    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)


def _extract_docx(data):
    try:
        import docx
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"DOCX support unavailable: {exc}")

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Could not read the DOCX file: {exc}")

    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Include table cell text too.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n\n".join(parts)


def _extract_pdf(data, lang=None):
    try:
        from pypdf import PdfReader
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"PDF support unavailable: {exc}")

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"Could not read the PDF file: {exc}")

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")

    text = "\n\n".join(p for p in pages if p.strip())

    # If the PDF has embedded text, we're done.
    if text.strip():
        return text

    # Otherwise it's likely a scanned/image-only PDF -> render each page to an
    # image with Poppler (pdf2image) and OCR it. Falls back to pypdf's embedded
    # images if Poppler isn't available.
    if not ocr_available():
        raise OCRUnavailable(
            "This looks like a scanned PDF, but text recognition (OCR) isn't "
            "available right now. Please upload a PDF that contains selectable text."
        )

    # Preferred path: pdf2image (needs the Poppler binary).
    try:
        from pdf2image import convert_from_bytes

        images = convert_from_bytes(data, dpi=300)
        from PIL import Image  # noqa: F401
        parts = []
        for pil_img in images:
            buf = io.BytesIO()
            pil_img.save(buf, format="PNG")
            parts.append(ocr_image_bytes(buf.getvalue(), lang=lang))
        ocr_text = "\n\n".join(p for p in parts if p.strip())
        if ocr_text.strip():
            return ocr_text
    except Exception:
        pass

    # Fallback: OCR any embedded page images pypdf can surface.
    try:
        blobs = []
        for page in reader.pages:
            for img in getattr(page, "images", []):
                blobs.append(img.data)
        ocr_text = "\n\n".join(ocr_image_bytes(b, lang=lang) for b in blobs)
        if ocr_text.strip():
            return ocr_text
    except Exception:
        pass

    return text


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(filename, data, lang=None):
    """Extract raw text from an uploaded file's bytes.

    ``filename`` is only used for its extension. Raises ExtractionError /
    OCRUnavailable on failure.
    """
    ext = os.path.splitext(filename or "")[1].lower()

    if ext in TEXT_EXTS:
        return _extract_txt(data)
    if ext in JSON_EXTS:
        return _extract_json(data)
    if ext in CSV_EXTS:
        return _extract_csv(data)
    if ext in EXCEL_EXTS:
        return _extract_excel(data)
    if ext in DOCX_EXTS:
        return _extract_docx(data)
    if ext in PDF_EXTS:
        return _extract_pdf(data, lang=lang)
    if ext in IMAGE_EXTS:
        return ocr_image_bytes(data, lang=lang)

    raise ExtractionError(f"Unsupported file type: {ext or 'unknown'}")


def split_paragraphs(text, min_words=3):
    """Split a text blob into clean paragraph strings.

    Splits on blank lines first; falls back to single newlines. Each returned
    paragraph is whitespace-collapsed and must contain at least ``min_words``.
    """
    if not text:
        return []

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Primary split on one-or-more blank lines.
    blocks = re.split(r"\n\s*\n", text)

    # If everything came back as one block, fall back to line-based splitting.
    if len(blocks) <= 1:
        blocks = text.split("\n")

    paragraphs = []
    for block in blocks:
        cleaned = re.sub(r"\s+", " ", block).strip()
        if cleaned and len(cleaned.split()) >= min_words:
            paragraphs.append(cleaned)

    return paragraphs
